from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "output/doc/Shotai_PRD_v1.1_完善版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "E8EEF7")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text))
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def set_doc_style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(31, 55, 99)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def build_doc():
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SHOTAI")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(31, 55, 99)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("AI 摄影助手 产品需求文档（PRD）").bold = True

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.add_run("v1.1 完善版 · 2026年6月")
    doc.add_paragraph()

    add_table(
        doc,
        ["字段", "内容"],
        [
            ["产品名称", "Shotai（写太）"],
            ["版本范围", "v1.0 MVP（网页版）+ v1.1 体验增强边界"],
            ["文档状态", "完善稿，待技术评审与隐私评审确认"],
            ["目标平台", "Web（桌面优先，移动端浏览器适配）"],
            ["建议技术栈", "React / TypeScript / API Proxy / Claude Vision API / Canvas 或 WebGL"],
            ["核心决策", "前端不直接暴露模型 API Key；AI 分析图片需经用户授权上传到模型服务"],
        ],
        widths=[1.6, 5.4],
    )

    add_heading(doc, "1. 产品概述", 1)
    add_heading(doc, "1.1 产品定位", 2)
    add_paragraph(
        doc,
        "Shotai 是一款面向摄影爱好者和轻量内容创作者的 AI 摄影助手，覆盖拍摄前规划与拍摄后调色两个场景，帮助用户从参考图中提取风格意图，并将其转化为可执行的拍摄建议和可调整的调色参数。"
    )
    add_heading(doc, "1.2 核心价值主张", 2)
    add_bullets(
        doc,
        [
            "拍前：上传目标风格照后，AI 解析场景、光线、构图和拍摄参数建议，降低试错成本。",
            "拍后：对比目标风格照与用户实拍照，AI 输出调色方向和结构化参数，用户可继续手动微调。",
            "可复用：用户可以保存一套调色参数为“风格方案”，后续应用到其他照片。",
            "可控隐私：本地预览、滑块处理和导出在浏览器内完成；需要 AI 分析时，明确告知图片会上传到模型服务。",
        ],
    )
    add_heading(doc, "1.3 目标用户", 2)
    add_table(
        doc,
        ["用户类型", "核心需求", "痛点"],
        [
            ["摄影爱好者", "想复刻参考图的光线、构图和色调", "懂审美但缺少系统摄影和修图知识"],
            ["内容创作者", "快速产出社交平台风格统一的图片", "批量图片保持一致风格成本高"],
            ["手机摄影用户", "希望获得简单易懂的拍摄和调色建议", "专业参数难理解，修图工具学习成本高"],
        ],
        widths=[1.5, 2.8, 2.7],
    )
    add_heading(doc, "1.4 MVP 范围原则", 2)
    add_bullets(
        doc,
        [
            "优先保证单张参考图到单张实拍图的调色闭环可用。",
            "批量处理和参数保存作为增强能力进入 v1.1，但 PRD 中提前定义数据结构和入口，避免后续返工。",
            "AI 输出只作为建议起点，用户最终通过滑块确认结果，避免承诺“自动调出完全同款”。",
            "移动端先保证上传、预览、分析、导出可用，不承诺重度批量处理性能。",
        ],
    )

    add_heading(doc, "2. 功能范围", 1)
    add_heading(doc, "2.1 功能模块总览", 2)
    add_table(
        doc,
        ["模块", "功能", "描述", "优先级"],
        [
            ["拍前分析", "单图上传", "支持 JPG、PNG、WEBP；HEIC 视浏览器能力提示兼容性", "P0"],
            ["拍前分析", "AI 场景分析", "识别光线、地点类型、构图、拍摄时机和可执行建议", "P0"],
            ["拍前分析", "拍摄参数建议", "输出光圈、快门、ISO、焦距范围，并提示适用前提", "P0"],
            ["拍后调色", "双图上传", "上传目标风格照与一张用户实拍照，进入调色分析", "P0"],
            ["拍后调色", "AI 调色分析", "输出文字解释和结构化调色参数", "P0"],
            ["拍后调色", "滑块实时微调", "亮度、对比度、饱和度、色温、阴影、高光", "P0"],
            ["拍后调色", "图片导出", "导出调色后的 JPG，文件名包含时间戳", "P0"],
            ["风格方案", "保存当前参数", "用户可保存当前六项参数为本地风格方案", "P1"],
            ["风格方案", "应用已保存参数", "选择保存方案并应用到当前图片", "P1"],
            ["风格预设", "内置预设库", "提供类似胶片模拟的风格预设，如富士感、日系清透、电影暖调等", "P1"],
            ["风格预设", "预设强度", "用户可调整预设应用强度，避免一键套用过重", "P1"],
            ["批量处理", "多张实拍图上传", "同一目标风格和同一参数批量应用到多张实拍图", "P1"],
            ["批量处理", "批量导出", "导出多张处理结果；移动端可降级为逐张导出", "P2"],
        ],
        widths=[1.1, 1.4, 3.7, 0.8],
    )

    add_heading(doc, "2.2 拍前分析详细需求", 2)
    add_heading(doc, "2.2.1 图片上传", 3)
    add_bullets(
        doc,
        [
            "支持点击选择文件和拖拽上传。",
            "MVP 支持 JPG、PNG、WEBP；HEIC 若浏览器无法解码，应提示用户转换格式。",
            "单张图片建议上限 10MB，硬性上限 20MB；超过建议上限时先进行本地压缩再分析。",
            "上传后展示预览、文件名、格式、原始尺寸和压缩后尺寸。",
            "若图片 EXIF 方向异常，预览和导出需自动校正方向。",
        ],
    )
    add_heading(doc, "2.2.2 AI 分析输出", 3)
    add_bullets(
        doc,
        [
            "场景与环境：时间段、天气特征、地点类型、主要光源。",
            "构图要点：拍摄角度、构图法则、主体位置、前景和背景关系。",
            "拍摄参数：光圈、快门、ISO、焦距范围，并标明手机/相机适用差异。",
            "执行建议：最佳拍摄时机、需要避开的风险、可尝试的变体角度。",
            "置信提示：当图片信息不足时，需明确说明推断不确定性。",
        ],
    )

    add_heading(doc, "2.3 拍后调色详细需求", 2)
    add_heading(doc, "2.3.1 调色参数", 3)
    add_table(
        doc,
        ["参数", "范围", "实现方式", "说明"],
        [
            ["brightness", "-100 ~ +100", "Canvas 像素处理或 CSS filter", "整体明暗"],
            ["contrast", "-100 ~ +100", "Canvas 像素处理或 CSS filter", "明暗反差"],
            ["saturation", "-100 ~ +100", "Canvas 像素处理", "色彩浓度"],
            ["temperature", "-100 ~ +100", "RGB 通道偏移或 WebGL shader", "冷暖色偏移"],
            ["shadows", "-100 ~ +100", "按亮度阈值处理暗部", "暗部细节提升或压暗"],
            ["highlights", "-100 ~ +100", "按亮度阈值处理亮部", "亮部恢复或增强"],
        ],
        widths=[1.3, 1.2, 2.4, 2.1],
    )
    add_heading(doc, "2.3.2 AI 输出 Schema", 3)
    add_paragraph(doc, "AI 接口返回应拆分为 explanation 与 adjustments 两部分，前端只解析 adjustments。")
    add_table(
        doc,
        ["字段", "类型", "必填", "校验规则"],
        [
            ["explanation", "string", "是", "中文说明，不参与滑块解析"],
            ["adjustments.brightness", "number", "是", "整数，范围 -100 到 100"],
            ["adjustments.contrast", "number", "是", "整数，范围 -100 到 100"],
            ["adjustments.saturation", "number", "是", "整数，范围 -100 到 100"],
            ["adjustments.temperature", "number", "是", "整数，范围 -100 到 100"],
            ["adjustments.shadows", "number", "是", "整数，范围 -100 到 100"],
            ["adjustments.highlights", "number", "是", "整数，范围 -100 到 100"],
            ["confidence", "number", "否", "0 到 1，用于提示推荐可信度"],
        ],
        widths=[2.1, 1.1, 0.8, 3.0],
    )
    add_bullets(
        doc,
        [
            "字段缺失时使用默认值 0，并向用户提示“部分参数未识别，已使用默认值”。",
            "字段越界时按 -100 到 100 截断。",
            "AI 返回无法解析时，保留文字分析，滑块保持当前值，并允许用户重新分析。",
            "同一张图片重复分析时，应覆盖上一轮 AI 推荐，但不覆盖用户手动保存的风格方案。",
            "AI 应只提取色彩、明暗、反差、颗粒感等风格特征，不以目标图和实拍图的内容相似度作为是否可用的判断条件。",
        ],
    )

    add_heading(doc, "2.4 风格方案保存", 2)
    add_paragraph(doc, "用于解决用户“想保留这套参数”的需求。MVP 可先做本地保存，后续账号体系再同步到云端。")
    add_table(
        doc,
        ["需求", "说明", "优先级"],
        [
            ["保存当前参数", "用户点击保存后输入方案名称，保存六项滑块参数、创建时间、来源图片摘要", "P1"],
            ["方案列表", "展示本地已保存方案，支持选择、重命名、删除", "P1"],
            ["应用方案", "将选中方案应用到当前实拍图，用户仍可继续微调", "P1"],
            ["导出方案", "导出 JSON 文件，便于备份或分享", "P2"],
            ["云端同步", "登录后跨设备保存风格方案", "V2.0"],
        ],
        widths=[1.6, 4.4, 1.0],
    )
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["presetId", "string", "本地生成 UUID"],
            ["name", "string", "用户自定义名称，默认“未命名风格”"],
            ["adjustments", "object", "六项调色参数"],
            ["presetType", "string", "builtIn 或 custom"],
            ["strength", "number", "预设强度，范围 0 到 100，默认 100"],
            ["source", "object", "可选：目标图尺寸、文件名摘要，不保存原图"],
            ["createdAt", "string", "ISO 时间"],
            ["updatedAt", "string", "ISO 时间"],
        ],
        widths=[1.6, 1.2, 4.2],
    )

    add_heading(doc, "2.5 风格预设与胶片模拟", 2)
    add_paragraph(
        doc,
        "Shotai 的拍后调色不要求目标图与实拍图内容一致。目标图或内置预设本质上提供“风格参考”，系统应抽象出色彩倾向、反差、明暗层次、饱和度、冷暖和颗粒感等可迁移参数。"
    )
    add_table(
        doc,
        ["预设类型", "示例", "说明", "优先级"],
        [
            ["内置胶片感", "富士感、经典负片、电影暖调", "不直接使用品牌官方 LUT 名称，避免商标和误导风险；文案可写“类似富士感”", "P1"],
            ["场景风格", "日系清透、港风夜景、冷调街拍、暖调人像", "便于普通用户按审美选择", "P1"],
            ["AI 提取风格", "从目标图生成一套参数", "将参考图转化为用户可保存的自定义预设", "P1"],
            ["用户自定义", "我的咖啡店风格、夏日海边", "用户保存当前滑块参数后形成个人预设", "P1"],
        ],
        widths=[1.4, 2.0, 2.8, 0.8],
    )
    add_bullets(
        doc,
        [
            "预设应作为调色起点，而非最终不可编辑滤镜；应用后六项滑块必须可见且可调整。",
            "内置预设需提供强度滑块，范围 0 到 100，默认 70 或 100 由设计评审确认。",
            "内置预设命名避免直接声称等同于 Fuji、Kodak、Leica 等品牌官方效果；可以使用“富士感”“胶片负片感”等描述性命名。",
            "AI 从参考图提取的预设应允许保存为自定义预设，并可应用到单张或多张实拍图。",
            "若未来支持 LUT、曲线或 HSL，需保留 presetVersion 字段，兼容旧版六项参数预设。",
        ],
    )

    add_heading(doc, "2.6 多张实拍图与批量处理", 2)
    add_paragraph(doc, "用于解决用户“想上传多张自己拍的照片”的需求。建议作为 v1.1 增强功能，不阻塞 MVP。")
    add_bullets(
        doc,
        [
            "用户在拍后调色中可上传一张目标风格照和多张实拍照。",
            "AI 分析默认只基于目标风格照和用户选中的第一张实拍照生成参数。",
            "生成参数后，用户可选择“应用到全部实拍图”。",
            "每张图允许独立微调，但默认复用同一套参数。",
            "桌面端最多一次处理 20 张；移动端最多一次处理 5 张，超过时提示分批上传。",
            "批量导出桌面端可打包 ZIP；移动端先支持逐张导出，ZIP 作为 P2。",
        ],
    )
    add_table(
        doc,
        ["极端情况", "产品处理方式"],
        [
            ["多张实拍图尺寸差异很大", "统一使用各自原始宽高处理，不强制裁切；预览容器等比缩放"],
            ["部分图片处理失败", "失败项单独标记，允许重试或跳过，不影响其他图片导出"],
            ["移动端内存不足", "限制并发处理数量，必要时降级为逐张处理"],
            ["用户想对每张图微调不同参数", "在批量列表中支持单张覆盖参数，显示“已自定义”状态"],
            ["用户更换目标风格照", "提示现有 AI 推荐可能不再适用，用户确认后重新分析"],
            ["用户中途取消批量处理", "停止待处理队列，保留已完成结果，允许继续处理剩余图片"],
            ["浏览器切后台或锁屏", "展示处理中提示，恢复前台后同步每张图状态；不保证后台持续处理"],
            ["导出文件名冲突", "自动使用原文件名 + preset 名称 + 时间戳生成唯一文件名"],
        ],
        widths=[2.4, 4.6],
    )

    add_heading(doc, "3. 用户流程", 1)
    add_heading(doc, "3.1 拍前分析流程", 2)
    add_numbered(
        doc,
        [
            "用户进入“拍前分析”标签页。",
            "上传目标风格照片，系统展示预览和基础信息。",
            "用户点击“分析拍摄方案”。",
            "系统如需上传到模型服务，展示授权说明；用户确认后继续。",
            "AI 返回结果后，系统展示场景分析、构图建议、参数建议和注意事项。",
            "用户可复制结果或重新上传另一张图片分析。",
        ],
    )
    add_heading(doc, "3.2 单张拍后调色流程", 2)
    add_numbered(
        doc,
        [
            "用户进入“拍后调色”标签页。",
            "分别上传目标风格照和一张用户实拍照。",
            "两张图片均可预览后，“分析调色方案”按钮激活。",
            "用户确认图片上传到模型服务后发起分析。",
            "AI 返回文字解释和结构化参数，系统自动设置滑块。",
            "用户拖动滑块微调，预览实时刷新。",
            "用户可保存参数为风格方案、应用内置预设，或导出处理后的 JPG。",
        ],
    )
    add_heading(doc, "3.3 多张实拍图流程（v1.1）", 2)
    add_numbered(
        doc,
        [
            "用户上传一张目标风格照和多张实拍图。",
            "系统默认选中第一张实拍图作为 AI 对比样本。",
            "AI 生成一套基础调色参数，或用户选择一个内置/自定义预设。",
            "用户点击“应用到全部”后，列表中的实拍图全部生成预览。",
            "用户可进入单张详情进行独立微调。",
            "用户选择导出全部或导出选中图片。",
        ],
    )

    add_heading(doc, "4. 非功能需求", 1)
    add_heading(doc, "4.1 性能", 2)
    add_table(
        doc,
        ["项目", "目标", "说明"],
        [
            ["首次加载", "< 3s", "桌面宽带环境，首屏可交互"],
            ["滑块响应", "< 100ms", "单张 3000px 长边以内图片"],
            ["AI 拍前分析", "< 10s", "不含用户确认时间"],
            ["AI 拍后分析", "< 15s", "双图上传与模型响应"],
            ["批量处理", "前台不卡死", "使用队列或 Worker，展示进度"],
        ],
        widths=[1.8, 1.4, 3.8],
    )
    add_heading(doc, "4.2 隐私与安全", 2)
    add_bullets(
        doc,
        [
            "前端不得内置 Claude API Key；模型调用必须通过受控 API Proxy 或用户自带 Key 模式。",
            "本地调色预览、滑块调整、参数保存和导出不上传服务器。",
            "使用 AI 分析时，图片会通过 HTTPS 发送到模型服务，界面需明确告知用户。",
            "服务端默认不持久化原图；若需要日志，仅记录请求 ID、耗时、错误码和匿名统计，不记录图片内容。",
            "本地保存风格方案默认不保存原图，仅保存参数和可选的非敏感摘要。",
        ],
    )
    add_heading(doc, "4.3 兼容性", 2)
    add_bullets(
        doc,
        [
            "桌面：Chrome 100+、Safari 15+、Firefox 100+、Edge 100+。",
            "移动端：iOS Safari 15+、Android Chrome 100+。",
            "最小宽度：375px。",
            "HEIC 支持依赖浏览器解码能力；不支持时必须给出明确提示。",
        ],
    )
    add_heading(doc, "4.4 可访问性", 2)
    add_bullets(
        doc,
        [
            "所有图标按钮提供 aria-label。",
            "核心操作支持键盘 Tab 导航。",
            "错误状态不能只依赖颜色表达。",
            "文本和控件颜色对比度满足 WCAG AA。",
        ],
    )

    add_heading(doc, "5. 异常与极端场景", 1)
    add_table(
        doc,
        ["场景", "处理策略", "优先级"],
        [
            ["上传文件格式不支持", "阻止上传并提示支持格式", "P0"],
            ["图片超过 20MB", "阻止分析，提示压缩或更换图片", "P0"],
            ["图片像素面积过大", "除 MB 限制外增加最大像素面积限制；本地生成分析用缩略图", "P0"],
            ["超长图、全景图或截图", "等比预览并提示该类图片不适合摄影风格分析；允许仅做手动调色", "P1"],
            ["EXIF 方向错误", "预览与导出自动校正", "P0"],
            ["色彩空间不一致", "统一转为 sRGB 预览和导出；提示 Display P3 等广色域图片可能存在色差", "P1"],
            ["低清晰度或重度压缩图片", "允许分析，但提示结果可信度较低", "P1"],
            ["透明 PNG 导出 JPG", "提示透明区域将填充白色或提供 PNG 导出选项", "P1"],
            ["AI 接口超时", "展示重试按钮，不清空已上传图片", "P0"],
            ["AI 返回无效 JSON", "保留文字说明，滑块不覆盖或使用默认值，允许重新分析", "P0"],
            ["AI 推荐参数过于极端", "前端将默认应用值限制在安全区间，用户可手动拉到完整范围", "P0"],
            ["同一图片多次分析结果不同", "保留最近一次分析结果，并提示 AI 推荐存在波动", "P1"],
            ["模型拒绝分析某些图片", "展示模型拒绝原因的通用说明，并允许用户改用手动调色或预设", "P0"],
            ["图片提示词注入", "当图片中包含文字指令或截图内容时，系统提示约束模型忽略这些指令，只分析摄影风格", "P0"],
            ["API 额度不足或服务异常", "展示明确错误，不进入无限加载", "P0"],
            ["用户刷新页面", "MVP 不恢复上传图片；已保存风格方案保留在本地", "P1"],
            ["用户上传多张实拍图", "v1.1 支持批量应用同一参数；MVP 则提示一次只处理一张", "P1"],
            ["用户想保留参数", "支持保存为本地风格方案，可重命名、删除、再次应用", "P1"],
            ["本地预设数量过多", "本地最多保存 100 套自定义预设，超出时提示删除旧方案或导出备份", "P1"],
            ["预设名称重复", "允许重复但自动展示创建时间；导出文件名自动加序号", "P1"],
            ["用户误删预设", "删除前二次确认；删除后短时间提供撤销入口", "P1"],
            ["旧版预设与新版参数不兼容", "通过 presetVersion 做迁移；无法迁移时保留原始 JSON 并提示部分参数不可用", "P2"],
            ["清除浏览器缓存导致预设丢失", "在本地保存说明中提示风险，并提供 JSON 导出备份", "P1"],
            ["用户上传敏感人物照片", "AI 分析前展示上传说明，提供取消入口", "P0"],
            ["用户上传第三方商业摄影图", "条款提示用户需拥有上传和处理图片的合法权利", "P1"],
        ],
        widths=[2.2, 4.0, 0.8],
    )

    add_heading(doc, "6. 数据模型", 1)
    add_heading(doc, "6.1 前端状态", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["beforeImage", "ImageState | null", "拍前分析上传图片"],
            ["targetImage", "ImageState | null", "拍后目标风格图"],
            ["userImages", "ImageState[]", "用户实拍图列表，MVP 长度为 1"],
            ["activeUserImageId", "string | null", "当前正在预览或微调的实拍图"],
            ["adjustments", "AdjustmentValues", "当前全局推荐参数"],
            ["perImageOverrides", "Record<string, AdjustmentValues>", "单张图片自定义参数"],
            ["presets", "Preset[]", "本地保存的风格方案"],
            ["activePresetId", "string | null", "当前应用的内置或自定义预设"],
            ["presetStrength", "number", "当前预设强度，范围 0 到 100"],
            ["analysisState", "idle | loading | success | error", "AI 分析状态"],
        ],
        widths=[2.0, 2.0, 3.0],
    )
    add_heading(doc, "6.2 本地存储", 2)
    add_bullets(
        doc,
        [
            "使用 localStorage 或 IndexedDB 保存风格方案。",
            "自定义预设建议使用 IndexedDB；localStorage 仅适合少量纯参数方案。",
            "MVP 不保存原图，避免本地存储膨胀和隐私风险。",
            "若用户导入/导出方案，使用 JSON 文件承载参数，不包含图片。",
            "本地预设最多保存 100 套，超过上限时引导用户删除或导出备份。",
        ],
    )

    add_heading(doc, "7. 埋点与成功指标", 1)
    add_table(
        doc,
        ["指标", "目标值", "埋点口径"],
        [
            ["上传成功率", "> 95%", "upload_success / upload_start"],
            ["AI 分析成功率", "> 90%", "analysis_success / analysis_start"],
            ["调色预览响应", "< 100ms", "滑块输入到 Canvas 刷新的客户端耗时"],
            ["导出成功率", "> 98%", "export_success / export_start"],
            ["用户完成率", "> 70%", "上传图片后完成分析或导出的 session 比例"],
            ["参数保存率", "观察指标", "save_preset_success / analysis_success"],
            ["预设应用率", "观察指标", "apply_preset_click / upload_success"],
            ["批量应用使用率", "观察指标", "batch_apply_click / multi_upload_success"],
        ],
        widths=[1.6, 1.2, 4.2],
    )

    add_heading(doc, "8. 迭代规划", 1)
    add_table(
        doc,
        ["版本", "范围", "说明"],
        [
            ["v1.0 MVP", "单图拍前分析、单张拍后调色、导出", "优先验证核心闭环"],
            ["v1.1", "内置预设、参数保存、多张实拍图、批量应用、单张覆盖参数", "解决风格复用和多图场景"],
            ["v1.2", "更多预设包、批量 ZIP 导出、方案 JSON 导入导出、移动端拍摄上传", "提升效率和分享能力"],
            ["v2.0", "账号系统、云端同步、社区风格库、Lightroom .xmp 导出", "平台化能力"],
        ],
        widths=[1.1, 2.8, 3.1],
    )

    add_heading(doc, "9. 开放问题", 1)
    add_table(
        doc,
        ["#", "问题", "建议决策 / 待确认"],
        [
            ["1", "是否必须完全无后端？", "不建议。若使用模型 API，应采用 API Proxy 或用户自带 Key。"],
            ["2", "图片上传到第三方 AI 服务是否可接受？", "需在产品界面明确告知，并确认隐私条款。"],
            ["3", "MVP 是否包含多张实拍图？", "建议不包含，v1.1 支持；但数据模型提前兼容数组。"],
            ["4", "参数保存是否进入 MVP？", "若目标是提高复用价值，建议至少提供本地保存作为 P1。"],
            ["5", "内置预设是否进入 MVP？", "建议至少提供 6 到 8 个基础风格预设作为 P1。"],
            ["6", "是否需要 RAW 格式？", "MVP 不支持，v1.2 后评估。"],
            ["7", "调色算法是否追求专业级准确？", "MVP 先做视觉近似，后续评估 WebGL、LUT、曲线和 HSL。"],
            ["8", "预设命名是否涉及品牌风险？", "避免使用官方商标式命名，采用描述性名称并做法务确认。"],
            ["9", "商标和域名是否可用？", "产品命名前需检索 Shotai / 写太 的商标、域名和应用商店冲突。"],
        ],
        widths=[0.5, 2.4, 4.1],
    )

    add_heading(doc, "10. MVP 验收清单", 1)
    add_bullets(
        doc,
        [
            "用户可以上传一张参考图并获得拍摄建议。",
            "用户可以上传目标风格图和一张实拍图并获得调色建议。",
            "AI 参数可自动填入六项滑块，且字段越界和缺失不会导致页面崩溃。",
            "用户可以手动调节滑块，并实时看到预览变化。",
            "用户可以导出处理后的 JPG。",
            "用户可以应用至少一套内置预设，并在应用后继续手动调整。",
            "API 超时、格式不支持、图片过大、JSON 解析失败都有明确提示。",
            "前端不暴露模型 API Key。",
            "界面明确说明 AI 分析需要上传图片到模型服务。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Shotai PRD v1.1 完善版 · 2026年6月"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
